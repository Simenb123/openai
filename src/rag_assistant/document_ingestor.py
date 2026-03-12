from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Tuple

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    metadata: Dict[str, Any]


class DocumentIngestor:
    """Leser .txt/.pdf/.docx og returnerer tekst + metadata.

    Design:
    - Best-effort: bruker tydelige feilmeldinger
    - Robust mot norske tegn
    """

    def parse_file(self, path: str | Path) -> ParsedDocument:
        p = Path(path)
        if not p.exists() or not p.is_file():
            raise FileNotFoundError(f"Finner ikke fil: {p}")

        ext = p.suffix.lower().lstrip(".")  # "txt" / "pdf" / "docx" / "" (ingen endelse)
        if ext == "":
            # Extensionless: vi antar tekst (praktisk for kilder hentet fra Lovdata o.l.)
            ext = "txt"

        if ext == "txt":
            text = self._read_text_file(p)
            meta = self._base_meta(p, ext="txt")
            return ParsedDocument(text=text, metadata=meta)

        if ext == "pdf":
            text, n_pages = self._read_pdf(p)
            meta = self._base_meta(p, ext="pdf")
            meta["n_pages"] = n_pages
            return ParsedDocument(text=text, metadata=meta)

        if ext == "docx":
            text = self._read_docx(p)
            meta = self._base_meta(p, ext="docx")
            return ParsedDocument(text=text, metadata=meta)

        raise ValueError(f"Ikke støttet filtype: .{ext}")

    def _base_meta(self, p: Path, *, ext: str) -> Dict[str, Any]:
        return {
            "source_path": str(p.resolve()),
            "file_name": p.name,
            "file_ext": ext,
        }

    def _read_text_file(self, p: Path) -> str:
        data = p.read_bytes()
        # Vanlige encoding-varianter (inkl. BOM)
        for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        # fallback: ignorer feil
        return data.decode("utf-8", errors="ignore")

    def _read_pdf(self, p: Path) -> Tuple[str, int]:
        try:
            from PyPDF2 import PdfReader  # lokal import for testbarhet
        except Exception as e:  # pragma: no cover
            raise RuntimeError("PyPDF2 er ikke installert. Kjør: pip install -r requirements.txt") from e

        reader = PdfReader(str(p))
        texts = []
        for page in reader.pages:
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            if t:
                texts.append(t)
        return "\n\n".join(texts).strip(), len(reader.pages)

    def _read_docx(self, p: Path) -> str:
        try:
            from docx import Document  # lokal import for testbarhet
        except Exception as e:  # pragma: no cover
            raise RuntimeError("python-docx er ikke installert. Kjør: pip install -r requirements.txt") from e

        doc = Document(str(p))
        parts = []
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)
        return "\n".join(parts).strip()
